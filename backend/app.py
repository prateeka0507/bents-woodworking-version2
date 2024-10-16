import os
import uuid
import re
import logging
from flask import Flask, render_template, request, jsonify, session
from werkzeug.utils import secure_filename
from docx import Document
from dotenv import load_dotenv
from langchain_openai import OpenAIEmbeddings, ChatOpenAI
from langchain_pinecone import PineconeVectorStore
from langchain.chains import ConversationalRetrievalChain
from langchain.schema import Document as LangchainDocument
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.prompts import ChatPromptTemplate, SystemMessagePromptTemplate, HumanMessagePromptTemplate
from pinecone import Pinecone, ServerlessSpec
import langsmith
from flask_cors import CORS
import psycopg2
from psycopg2.extras import RealDictCursor
import base64
from tenacity import retry, stop_after_attempt, wait_fixed, retry_if_exception_type
import time

class LLMResponseError(Exception):
    pass

class LLMResponseCutOff(LLMResponseError):
    pass

class LLMNoResponseError(LLMResponseError):
    pass

load_dotenv()

app = Flask(__name__)
CORS(app, resources={r"/*": {"origins": [
    "http://localhost:5002",
    "http://localhost:5173"
]}})

app.secret_key = os.urandom(24)  # Set a secret key for sessions


# Access your API keys (set these in environment variables)
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
PINECONE_API_KEY = os.getenv("PINECONE_API_KEY")
TRANSCRIPT_INDEX_NAMES = ["bents", "shop-improvement", "tool-recommendations"]
PRODUCT_INDEX_NAME = "bents-woodworking-products"
os.environ["LANGCHAIN_TRACING_V2"] = "true"
os.environ["LANGCHAIN_API_KEY"] = os.getenv("LANGSMITH_API_KEY")
os.environ["LANGCHAIN_ENDPOINT"] = "https://api.smith.langchain.com"
os.environ["LANGCHAIN_PROJECT"] = "jason-json"


# Video title list
VIDEO_TITLE_LIST = [
    "5 Modifications I Made In My Garage Shop - New Shop Part 5",
    "2020 Shop Tour",
    "American Green Lights",
    "Assembly Table and Miter Saw Station",
    "Complete Mr Cool Install",
    "Every track saw owner could use this",
    "How To Install Mr Cool DIY Series",
    "I Built a Wall in my Garage",
    "Moving A Woodworking Shop - New Shop Part 2",
    "My shop is soundproof",
    "People Told Me My Garage Door Would Explode",
    "The biggest advancement in dust collection",
    "Using SketchUp To Design Woodworking Shop - New Shop Part 1",
    "8 Tools I Regret Not Buying Sooner",
    "10 Tools Every Woodworker Should Own",
    "10 woodworking tools I regret not buying sooner",
    "10 Woodworking tools you will not regret",
    "11 woodworking tools you need to own",
    "12 Tools I will Never REGRET Buying",
    "15 cabinet tools I do not regret",
    "15 Woodworking Tools You Will not Regret",
    "25 tools I regret not buying sooner",
    "Every track saw owner could use this",
    "FINALLY! The sprayer I have been waiting for",
    "I would not buy these with your money",
    "Stop wasting your money on the wrong ones",
    "The 5 TSO tools you cannot live without",
    "Track Saw Square Comparison TSO ProductsBench Dogs UKWoodpeckers ToolsInsta Rail Square"
]

# Initialize Langchain components
embeddings = OpenAIEmbeddings(openai_api_key=OPENAI_API_KEY)
llm = ChatOpenAI(openai_api_key=OPENAI_API_KEY, model="gpt-4o-mini", temperature=0)

# Initialize Pinecone
pc = Pinecone(api_key=PINECONE_API_KEY)

# Create or connect to the Pinecone indexes
for INDEX_NAME in TRANSCRIPT_INDEX_NAMES + [PRODUCT_INDEX_NAME]:
    if INDEX_NAME not in pc.list_indexes().names():
        pc.create_index(
            name=INDEX_NAME,
            dimension=1536,  # OpenAI embeddings dimension
            metric='cosine',
            spec=ServerlessSpec(cloud='aws', region='us-east-1')
        )

# Create VectorStores
transcript_vector_stores = {name: PineconeVectorStore(index=pc.Index(name), embedding=embeddings, text_key="text") for name in TRANSCRIPT_INDEX_NAMES}
product_vector_store = PineconeVectorStore(index=pc.Index(PRODUCT_INDEX_NAME), embedding=embeddings, text_key="tags")

# System instructions
SYSTEM_INSTRUCTIONS = """You are an AI assistant specialized in information retrieval from text documents.
        Always provide your responses in English, regardless of the language of the input or context.
        When given a document and a query:
        1. Analyze the document content and create an efficient index of key terms, concepts, and their locations within the text.
        2. When a query is received, use the index to quickly locate relevant sections of the document.
        3. Extract the most relevant information from those sections to form a concise and accurate answer.
        4. Always include the exact relevant content from the document, starting from the beginning of the relevant section. Use quotation marks to denote direct quotes.
        5. If applicable, provide a timestamp or location reference for where the information was found in the original document.
        6. After providing the direct quote, summarize or explain the answer if necessary.
        7. If the query cannot be answered from the given document, state this clearly.
        8. Always prioritize accuracy over speed. If you're not certain about an answer, say so.
        9. For multi-part queries, address each part separately and clearly.
        10. Aim to provide responses within seconds, even for large documents.
        11. please only Provide the timestamp for where the information was found in the original video. must Use the format {{timestamp:MM:SS}} for timestamps under an hour, and {{timestamp:HH:MM:SS}} for longer videos.
        12. Do not include any URLs in your response. Just provide the timestamps in the specified format.
        13. When referencing timestamps that may be inaccurate, you can use language like "around", "approximately", or "in the vicinity of" to indicate that the exact moment may vary slightly.
        Remember, always respond in English, even if the query or context is in another language.
        Always represent the speaker as Jason bent.You are an assistant expert representing Jason Bent as jason bent on woodworking response. Answer questions based on the provided context. The context includes timestamps in the format [Timestamp: HH:MM:SS]. When referencing information, include these timestamps in the format {{timestamp:HH:MM:SS}}.
Then show that is in generated response with the provided context.
"""

logging.basicConfig(level=logging.DEBUG)

def get_matched_products(video_title):
    logging.debug(f"Attempting to get matched products for title: {video_title}")
    try:
        conn = psycopg2.connect(os.getenv("POSTGRES_URL"))
        with conn.cursor(cursor_factory=RealDictCursor) as cur:
            # Query for partial matches in the title, case-insensitive
            query = """
                SELECT * FROM products 
                WHERE LOWER(tags) LIKE LOWER(%s)
            """
            search_term = f"%{video_title}%"
            logging.debug(f"Executing SQL query: {query} with search term: {search_term}")
            cur.execute(query, (search_term,))
            matched_products = cur.fetchall()
            logging.debug(f"Raw matched products from database: {matched_products}")
        conn.close()

        # Process the results
        related_products = [
            {
                'id': product['id'],
                'title': product['title'],
                'tags': product['tags'].split(',') if product['tags'] else [],
                'link': product['link'],
                'image_data': product['image_data'] if 'image_data' in product else None
            } for product in matched_products
        ]

        logging.debug(f"Processed related products: {related_products}")
        return related_products

    except Exception as e:
        logging.error(f"Error in get_matched_products: {str(e)}", exc_info=True)
        return []

def verify_database():
    try:
        conn = psycopg2.connect(os.getenv("POSTGRES_URL"))
        with conn.cursor(cursor_factory=RealDictCursor) as cur:
            cur.execute("SELECT COUNT(*) FROM products")
            count = cur.fetchone()['count']
            logging.info(f"Total products in database: {count}")
            
            cur.execute("SELECT title FROM products LIMIT 5")
            sample_titles = [row['title'] for row in cur.fetchall()]
            logging.info(f"Sample product titles: {sample_titles}")
        conn.close()
        return True
    except Exception as e:
        logging.error(f"Database verification failed: {str(e)}", exc_info=True)
        return False

def process_answer(answer, urls):
    def replace_timestamp(match):
        timestamp = match.group(1)
        full_urls = [combine_url_and_timestamp(url, timestamp) for url in urls if url]
        return f"[video]({','.join(full_urls)})"
    
    processed_answer = re.sub(r'\{timestamp:([^\}]+)\}', replace_timestamp, answer)
    
    video_links = re.findall(r'\[video\]\(([^\)]+)\)', processed_answer)
    video_dict = {f'[video{i}]': link.split(',') for i, link in enumerate(video_links)}
    
    for i, (placeholder, links) in enumerate(video_dict.items()):
        processed_answer = processed_answer.replace(f'[video]({",".join(links)})', placeholder)
    
    return processed_answer, video_dict

def combine_url_and_timestamp(base_url, timestamp):
    parts = timestamp.split(':')
    if len(parts) == 2:
        minutes, seconds = map(int, parts)
        total_seconds = minutes * 60 + seconds
    elif len(parts) == 3:
        hours, minutes, seconds = map(int, parts)
        total_seconds = hours * 3600 + minutes * 60 + seconds
    else:
        raise ValueError("Invalid timestamp format")

    if '?' in base_url:
        return f"{base_url}&t={total_seconds}"
    else:
        return f"{base_url}?t={total_seconds}"

def extract_text_from_docx(file):
    doc = Document(file)
    text = "\n".join([para.text for para in doc.paragraphs])
    return text

def extract_metadata_from_text(text):
    title = text.split('\n')[0] if text else "Untitled Video"
    return {"title": title}

def upsert_transcript(transcript_text, metadata, index_name):
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
    chunks = text_splitter.split_text(transcript_text)
    
    documents = []
    for i, chunk in enumerate(chunks):
        chunk_metadata = metadata.copy()
        chunk_metadata['chunk_id'] = f"{metadata['title']}_chunk_{i}"
        chunk_metadata['url'] = metadata.get('url', '')
        documents.append(LangchainDocument(page_content=chunk, metadata=chunk_metadata))
    
    transcript_vector_stores[index_name].add_documents(documents)

@app.route('/')
@app.route('/database')
def serve_spa():
    return render_template('index.html')

@app.route('/chat', methods=['POST'])
def chat():
    try:
        data = request.json
        logging.debug(f"Received data: {data}")

        user_query = data['message'].strip()
        selected_index = data['selected_index']
        chat_history = data.get('chat_history', [])

        logging.debug(f"Chat history received: {chat_history}")

        # Initial input validation
        if not user_query or user_query in ['.', ',', '?', '!']:
            return jsonify({
                'response': "I'm sorry, but I didn't receive a valid question. Could you please ask a complete question?",
                'related_products': [],
                'urls': [],
                'contexts': [],
                'video_links': {}
            })

        # Format chat history for ConversationalRetrievalChain
        formatted_history = []
        for i in range(0, len(chat_history) - 1, 2):
            human = chat_history[i]
            ai = chat_history[i + 1] if i + 1 < len(chat_history) else ""
            formatted_history.append((human, ai))

        logging.debug(f"Formatted chat history: {formatted_history}")

        # Relevance check
        relevance_check_prompt = f"""
        Given the following question or message and the chat history, determine if it is:
        1. A greeting or general conversation starter
        2. Related to woodworking, tools, home improvement, or the assistant's capabilities and also query about bents-woodworking youtube channel general questions.
        3. Related to the company, its products, services, or business operations
        4. A continuation or follow-up question to the previous conversation
        5. Related to violence, harmful activities, or other inappropriate content
        6. Completely unrelated to the above topics and not a continuation of the conversation
        7. if user is asking about jason bents.

        If it falls under category 1, respond with 'GREETING'.
        If it falls under categories 2, 3, 4 or 7 respond with 'RELEVANT'.
        If it falls under category 5, respond with 'INAPPROPRIATE'.
        If it falls under category 6, respond with 'NOT RELEVANT'.

        Chat History:
        {formatted_history[-3:] if formatted_history else "No previous context"}

        Current Question: {user_query}
        
        Response (GREETING, RELEVANT, INAPPROPRIATE, or NOT RELEVANT):
        """
        
        relevance_response = llm.predict(relevance_check_prompt)
        
        if "GREETING" in relevance_response.upper():
            greeting_response = llm.predict("Generate a friendly greeting response for a woodworking assistant.")
            return jsonify({
                'response': greeting_response,
                'related_products': [],
                'urls': [],
                'contexts': [],
                'video_links': {}
            })
        elif "INAPPROPRIATE" in relevance_response.upper():
            return jsonify({
                'response': "I'm sorry, but this is outside my context of answering. Is there something else I can help you with regarding woodworking, tools, or home improvement?",
                'related_products': [],
                'urls': [],
                'contexts': [],
                'video_links': {}
            })
        elif "NOT RELEVANT" in relevance_response.upper():
            return jsonify({
                'response': "I'm sorry, but I'm specialized in topics related to our company, woodworking, tools, and home improvement. I can also engage in general conversation or continue our previous discussion. Could you please ask a question related to these topics, continue our previous conversation, or start with a greeting?",
                'related_products': [],
                'urls': [],
                'contexts': [],
                'video_links': {}
            })

        # If we reach here, the query is relevant and not a greeting
        retriever = transcript_vector_stores[selected_index].as_retriever(search_kwargs={"k": 5})
        
        prompt = ChatPromptTemplate.from_messages([
            SystemMessagePromptTemplate.from_template(SYSTEM_INSTRUCTIONS),
            HumanMessagePromptTemplate.from_template("Context: {context}\n\nChat History: {chat_history}\n\nQuestion: {question}")
        ])
        
        qa_chain = ConversationalRetrievalChain.from_llm(
            llm=llm,
            retriever=retriever,
            combine_docs_chain_kwargs={"prompt": prompt},
            return_source_documents=True
        )
        
        try:
            result = retry_llm_call(qa_chain, user_query, formatted_history)
        except LLMResponseError as e:
            error_message = "Failed to get a complete response from the AI after multiple attempts."
            if isinstance(e, LLMNoResponseError):
                error_message = "The AI failed to generate a response after multiple attempts."
            return jsonify({'error': error_message}), 500
        except Exception as e:
            logging.error(f"Unexpected error in LLM call: {str(e)}")
            return jsonify({'error': 'An unexpected error occurred while processing your request.'}), 500
        
        initial_answer = result['answer']
        contexts = [doc.page_content for doc in result['source_documents']]
        source_documents = result['source_documents']

        # Extract video titles and URLs from all source documents
        video_titles = []
        urls = []
        for doc in source_documents:
            metadata = doc.metadata
            video_titles.append(metadata.get('title', "Unknown Video"))
            urls.append(metadata.get('url', None))

        logging.debug(f"Extracted video titles: {video_titles}")
        logging.debug(f"Extracted URLs: {urls}")

        processed_answer, video_dict = process_answer(initial_answer, urls)
        logging.debug(f"Processed answer: {processed_answer}")

        related_products = get_matched_products(video_titles[0] if video_titles else "Unknown Video")
        logging.debug(f"Retrieved matched products: {related_products}")

        response_data = {
            'response': processed_answer,
            'initial_answer': initial_answer,
            'related_products': related_products,
            'urls': urls,
            'contexts': contexts,
            'video_links': video_dict,
            'video_titles': video_titles
        }

        logging.debug(f"Response data: {response_data}")

        return jsonify(response_data)
    except Exception as e:
        logging.error(f"Error in chat route: {str(e)}", exc_info=True)
        return jsonify({'error': 'An error occurred processing your request'}), 500

@app.route('/api/user/<user_id>', methods=['GET'])
def get_user_data(user_id):
    try:
        # In a real application, you'd fetch this data from a database
        # For now, we'll return a dummy structure
        user_data = {
            'conversationsBySection': {
                "bents": [],
                "shop-improvement": [],
                "tool-recommendations": []
            },
            'searchHistory': [],
            'selectedIndex': "bents"
        }
        return jsonify(user_data)
    except Exception as e:
        logging.error(f"Error fetching user data: {str(e)}", exc_info=True)
        return jsonify({'error': 'An error occurred fetching user data'}), 500

@app.route('/upload_document', methods=['POST'])
def upload_document():
    if 'file' not in request.files:
        return jsonify({'success': False, 'message': 'No file part'})
    file = request.files['file']
    if file.filename == '':
        return jsonify({'success': False, 'message': 'No selected file'})
    
    index_name = request.form.get('index_name')
    if index_name not in TRANSCRIPT_INDEX_NAMES:
        return jsonify({'success': False, 'message': 'Invalid index name'})
    
    if file and file.filename.endswith('.docx'):
        filename = secure_filename(file.filename)
        file_path = os.path.join('/tmp', filename)
        file.save(file_path)
        
        transcript_text = extract_text_from_docx(file_path)
        metadata = extract_metadata_from_text(transcript_text)
        upsert_transcript(transcript_text, metadata, index_name)
        
        os.remove(file_path)
        return jsonify({'success': True, 'message': 'File uploaded and processed successfully'})
    else:
        return jsonify({'success': False, 'message': 'Invalid file format'})

@app.route('/documents')
def get_documents():
    try:
        conn = psycopg2.connect(os.getenv("POSTGRES_URL"))
        with conn.cursor(cursor_factory=RealDictCursor) as cur:
            cur.execute("SELECT * FROM products")
            documents = cur.fetchall()
        conn.close()
        return jsonify(documents)
    except Exception as e:
        print(f"Error in get_documents: {str(e)}")
        return jsonify({"error": str(e)}), 500

@app.route('/add_document', methods=['POST'])
def add_document():
    data = request.json
    try:
        conn = psycopg2.connect(os.getenv("POSTGRES_URL"))
        with conn.cursor(cursor_factory=RealDictCursor) as cur:
            cur.execute(
                "INSERT INTO products (title, tags, link) VALUES (%s, %s, %s) RETURNING id",
                (data['title'], ','.join(data['tags']), data['link'])
            )
            product_id = cur.fetchone()['id']
        conn.commit()
        conn.close()
        return jsonify({'success': True, 'product_id': product_id})
    except Exception as e:
        print(f"Error in add_document: {str(e)}")
        return jsonify({"error": str(e)}), 500

@app.route('/delete_document', methods=['POST'])
def delete_document():
    data = request.json
    try:
        conn = psycopg2.connect(os.getenv("POSTGRES_URL"))
        with conn.cursor() as cur:
            cur.execute("DELETE FROM products WHERE id = %s", (data['id'],))
        conn.commit()
        conn.close()
        return jsonify({'success': True})
    except Exception as e:
        print(f"Error in delete_document: {str(e)}")
        return jsonify({"error": str(e)}), 500

@app.route('/update_document', methods=['POST'])
def update_document():
    data = request.json
    try:
        conn = psycopg2.connect(os.getenv("POSTGRES_URL"))
        with conn.cursor() as cur:
            cur.execute(
                "UPDATE products SET title = %s, tags = %s, link = %s WHERE id = %s",
                (data['title'], ','.join(data['tags']), data['link'], data['id'])
            )
        conn.commit()
        conn.close()
        return jsonify({'success': True})
    except Exception as e:
        print(f"Error in update_document: {str(e)}")
        return jsonify({"error": str(e)}), 500

@retry(stop=stop_after_attempt(3), wait=wait_fixed(2), retry=retry_if_exception_type(LLMResponseError))
def retry_llm_call(qa_chain, query, chat_history):
    try:
        result = qa_chain({"question": query, "chat_history": chat_history})
        
        if result is None or 'answer' not in result or not result['answer']:
            raise LLMNoResponseError("LLM failed to generate a response")
        
        if result['answer'].endswith('...') or len(result['answer']) < 20:
            raise LLMResponseCutOff("LLM response appears to be cut off")
        
        return result
    except Exception as e:
        if isinstance(e, LLMResponseError):
            logging.error(f"LLM call failed: {str(e)}")
            raise
        logging.error(f"Unexpected error in LLM call: {str(e)}")
        raise LLMNoResponseError("LLM failed due to an unexpected error")

if __name__ == '__main__':
    verify_database()
    app.run(debug=True, port=5000)
